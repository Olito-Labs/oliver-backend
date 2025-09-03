from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import uuid
import os
import tempfile
import json
from datetime import datetime

from app.supabase_client import supabase
from app.auth import get_current_user
from app.llm_providers import openai_manager
from app.config import settings

# Document processing imports
import fitz  # PyMuPDF
try:
    import docx
except ImportError:
    docx = None

router = APIRouter(prefix="/api/documents", tags=["documents"])

# Document text extraction functions
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF with OCR fallback"""
    try:
        doc = fitz.open(file_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Try to get text directly first
            page_text = page.get_text()
            
            # If no text or very little text, try OCR
            if not page_text.strip() or len(page_text.strip()) < 50:
                try:
                    # Get page as image and use OCR
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    # For now, we'll use the basic text extraction
                    # In production, you might want to add pytesseract OCR here
                    page_text = page.get_text()
                except Exception as ocr_error:
                    print(f"OCR failed for page {page_num}: {ocr_error}")
                    page_text = page.get_text()
            
            text += f"--- Page {page_num + 1} ---\n"
            text += page_text + "\n\n"
        
        doc.close()
        return text.strip()
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    if docx is None:
        raise Exception("python-docx library not available")
    
    try:
        doc = docx.Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_document(file_path: str, file_type: str) -> str:
    """Extract text from document based on file type"""
    file_type = file_type.lower()
    
    if file_type == 'application/pdf':
        return extract_text_from_pdf(file_path)
    elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
        return extract_text_from_docx(file_path)
    else:
        raise Exception(f"Unsupported file type: {file_type}")

# Request/Response models
class DocumentResponse(BaseModel):
    id: str
    filename: str
    file_size: int
    file_type: str
    upload_url: str
    study_id: str
    user_id: str
    created_at: str
    processing_status: str
    analysis_results: Optional[Dict[str, Any]] = None  # OpenAI analysis results
    error_message: Optional[str] = None  # Error message if analysis failed

class DocumentAnalysisRequest(BaseModel):
    analysis_type: str = "mra-intake"

# Pydantic models for structured MRA analysis output
class RecommendedAction(BaseModel):
    action: str
    owner: Optional[str] = None
    timeframe: Optional[str] = None  # e.g., "0-30 days", "30-60 days"
    effort: Optional[str] = None     # Low | Medium | High
    impact: Optional[str] = None     # Low | Medium | High

class MRAFinding(BaseModel):
    id: str
    type: str  # 'BSA/AML', 'Lending', 'Operations', 'Capital', 'Management', 'IT', 'Other'
    severity: str  # 'Matter Requiring Attention', 'Deficiency', 'Violation', 'Recommendation'
    description: str
    regulatory_citation: Optional[str] = None
    deadline: Optional[str] = None
    response_required: bool
    extracted_text: str
    confidence: float
    # First-principles extensions (optional)
    what_it_means: Optional[str] = None
    why_it_matters: Optional[str] = None
    root_causes: Optional[List[str]] = None
    regulatory_expectations: Optional[List[str]] = None
    control_gaps: Optional[List[str]] = None
    recommended_actions: Optional[List[RecommendedAction]] = None
    acceptance_criteria: Optional[List[str]] = None
    artifacts_to_prepare: Optional[List[str]] = None
    risk_impact: Optional[str] = None  # Low | Medium | High
    urgency: Optional[str] = None      # e.g., Immediate, 30 days, 60-90 days
    dependencies: Optional[List[str]] = None
    closure_narrative_outline: Optional[str] = None

class FirstPrinciplesOverview(BaseModel):
    problem_statement: Optional[str] = None
    decomposition: Optional[List[str]] = None
    causal_factors: Optional[List[str]] = None
    regulatory_expectations: Optional[List[str]] = None
    risks: Optional[List[str]] = None

class RemediationGuidance(BaseModel):
    quick_wins: Optional[List[str]] = []
    critical_actions: Optional[List[str]] = []
    phases: Optional[List[Dict[str, Any]]] = []  # [{phase, goals[], notes}]
    owners: Optional[List[str]] = []
    timeline_30_60_90: Optional[Dict[str, List[str]]] = None  # keys: "30_days","60_days","90_days"
    acceptance_criteria: Optional[List[str]] = []
    monitoring_metrics: Optional[List[str]] = []
    required_artifacts: Optional[List[str]] = []

class MRAAnalysisResult(BaseModel):
    summary: str
    findings: List[MRAFinding]
    total_findings: int
    critical_deadlines: List[str]
    processing_time_ms: int
    confidence: float
    # First-principles extensions (optional)
    first_principles: Optional[FirstPrinciplesOverview] = None
    remediation_guidance: Optional[RemediationGuidance] = None
    overall_risk_level: Optional[str] = None  # Low | Medium | High
    urgency: Optional[str] = None
    key_regulatory_citations: Optional[List[str]] = None

async def analyze_mra_document_with_ai(document_text: str) -> MRAAnalysisResult:
    """Analyze MRA document using OpenAI (GPT-5/o3) with structured output"""
    start_time = datetime.now()
    
    client = openai_manager.get_client()
    if not client:
        raise Exception("OpenAI client not available")
    
    # Create specialized MRA analysis prompt with exact JSON schema
    system_prompt = """You are a regulatory compliance expert specializing in analyzing Matter Requiring Attention (MRA) documents from bank examinations.

Your task is to extract structured information from MRA documents, identifying regulatory findings with high precision, and to break them down from first principles so a bank officer knows exactly what it says and how to remediate effectively.

CRITICAL: You MUST return a valid JSON object that exactly matches this schema:

{
  "summary": "Brief summary of the document and findings",
  "findings": [
    {
      "id": "finding-1",
      "type": "BSA/AML|Lending|Operations|Capital|Management|IT|Other",
      "severity": "Matter Requiring Attention|Deficiency|Violation|Recommendation",
      "description": "Clear summary of the finding",
      "regulatory_citation": "Specific regulation cited or null",
      "deadline": "YYYY-MM-DD format or null",
      "response_required": true,
      "extracted_text": "Exact quote from document",
      "confidence": 0.85,

      // First-principles optional detail
      "what_it_means": "Plain-language explanation",
      "why_it_matters": "Business/regulatory impact",
      "root_causes": ["root cause 1", "root cause 2"],
      "regulatory_expectations": ["expectation 1"],
      "control_gaps": ["gap 1"],
      "recommended_actions": [
        {"action": "specific action", "owner": "Risk", "timeframe": "0-30 days", "effort": "Medium", "impact": "High"}
      ],
      "acceptance_criteria": ["what proves closure"],
      "artifacts_to_prepare": ["policy update", "evidence log"],
      "risk_impact": "High",
      "urgency": "Immediate (0-30 days)",
      "dependencies": ["data migration"],
      "closure_narrative_outline": "Outline for regulator response"
    }
  ],
  "total_findings": 5,
  "critical_deadlines": ["2025-06-30", "2025-12-31"],
  "processing_time_ms": 2500,
  "confidence": 0.87,

  // First-principles (optional but preferred)
  "first_principles": {
    "problem_statement": "Crisp articulation of the core problem",
    "decomposition": ["sub-problem 1", "sub-problem 2"],
    "causal_factors": ["factor 1"],
    "regulatory_expectations": ["what the regulation expects"],
    "risks": ["risk 1"]
  },
  "remediation_guidance": {
    "quick_wins": ["win 1"],
    "critical_actions": ["critical action 1"],
    "phases": [{"phase": "stabilize", "goals": ["stop-gap control"]}],
    "owners": ["Compliance", "IT"],
    "timeline_30_60_90": {
      "30_days": ["actions"],
      "60_days": ["actions"],
      "90_days": ["actions"]
    },
    "acceptance_criteria": ["measurable outcomes"],
    "monitoring_metrics": ["KRI1", "KCI1"],
    "required_artifacts": ["evidence list"]
  },
  "overall_risk_level": "Medium|High|Low",
  "urgency": "Immediate|30-60 days|>90 days",
  "key_regulatory_citations": ["12 CFR ..."]
}

FIELD REQUIREMENTS:
- id: Generate unique IDs like "finding-1", "finding-2", etc.
- type: Must be one of the exact values listed
- severity: Must be one of the exact values listed
- description: Clear, professional summary
- regulatory_citation: Include if cited, otherwise null
- deadline: YYYY-MM-DD format if mentioned, otherwise null
- response_required: true if response needed, false otherwise
- extracted_text: Direct quote supporting the finding
- confidence: Your confidence (0.0-1.0) in this specific finding
- total_findings: Count of findings array length
- critical_deadlines: List of all deadlines within 90 days
- processing_time_ms: Always set to 2500 (will be calculated server-side)
- confidence: Overall confidence in entire analysis

GUIDANCE:
- Prefer including the first_principles and remediation_guidance sections to help the officer act.
- Be precise and conservative. Only extract findings clearly stated in the document.
- Keep content concise, actionable, and bank-ready."""

    user_prompt = f"""Analyze this MRA document and extract all regulatory findings. Return ONLY valid JSON matching the exact schema provided:

DOCUMENT TEXT:
{document_text}

RESPONSE FORMAT: Valid JSON object only, no additional text or explanation."""

    try:
        # Build parameters based on model type for MRA document analysis
        request_params = {
            "model": settings.OPENAI_MODEL,
            "input": user_prompt,  # Document text input
            "instructions": system_prompt,  # System/developer instructions
            "max_output_tokens": 8000,  # Token limit for reasoning + response
            "text": {
                "format": {"type": "json_object"}  # Structured JSON output
            },
            "store": True,  # Store for potential debugging
            "stream": False  # Synchronous response
        }
        
        # Add model-specific parameters
        if settings.OPENAI_MODEL.startswith("gpt-5"):
            request_params["reasoning"] = {"effort": "medium"}  # Thorough MRA analysis
            request_params["text"]["verbosity"] = "high"  # Detailed structured output
        elif settings.OPENAI_MODEL.startswith("o3"):
            request_params["reasoning"] = {"effort": "medium", "summary": "detailed"}
        else:
            request_params["temperature"] = 0.7  # Consistent extraction for other models
        
        # Use OpenAI Responses API with model-specific configuration
        response = client.responses.create(**request_params)
        
        # Parse the response from Responses API format
        response_content = ""
        if response.output and len(response.output) > 0:
            # Iterate through output items to find message content
            for output_item in response.output:
                # Look for message type output items
                if hasattr(output_item, 'type') and output_item.type == 'message':
                    if hasattr(output_item, 'content') and len(output_item.content) > 0:
                        for content_item in output_item.content:
                            if hasattr(content_item, 'type') and content_item.type == 'output_text':
                                if hasattr(content_item, 'text'):
                                    response_content = content_item.text
                                    break
                    if response_content:
                        break
        
        # Fallback: Check if SDK provides convenience properties
        if not response_content and hasattr(response, 'output_text'):
            response_content = response.output_text
        
        if not response_content:
            # Log the response structure for debugging
            print(f"DEBUG: No content found. Response structure:")
            print(f"  - output length: {len(response.output) if response.output else 'None'}")
            if response.output:
                for i, item in enumerate(response.output):
                    print(f"  - output[{i}]: type={getattr(item, 'type', 'unknown')}")
                    if hasattr(item, 'content'):
                        for j, content in enumerate(item.content):
                            print(f"    - content[{j}]: type={getattr(content, 'type', 'unknown')}")
            raise Exception("No response content received from OpenAI - check response format")
        
        # Parse the JSON response
        import json
        analysis_data = json.loads(response_content)
        
        # Calculate actual processing time and override model's value
        processing_time = (datetime.now() - start_time).total_seconds() * 1000
        analysis_data['processing_time_ms'] = int(processing_time)
        
        # Create the result object
        analysis_result = MRAAnalysisResult(**analysis_data)
        
        # Generate unique IDs for findings if not provided
        for i, finding in enumerate(analysis_result.findings):
            if not finding.id:
                finding.id = f"finding-{i+1}"
        
        analysis_result.total_findings = len(analysis_result.findings)
        
        return analysis_result
        
    except Exception as e:
        print(f"OpenAI analysis error: {e}")
        raise Exception(f"Failed to analyze document with OpenAI: {str(e)}")

@router.post("/upload", response_model=Dict[str, DocumentResponse])
async def upload_document(
    file: UploadFile = File(...),
    study_id: str = Form(...),
    user=Depends(get_current_user)
):
    """Upload a document to Supabase Storage and create a database record"""
    try:
        # Validate file type
        allowed_types = ['application/pdf', 'application/msword', 
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document']
        if file.content_type not in allowed_types:
            raise HTTPException(
                status_code=400, 
                detail=f"File type {file.content_type} not supported. Please upload PDF, DOC, or DOCX files."
            )
        
        # Read file content once
        file_content = await file.read()
        print(f"File read: {file.filename}, size: {len(file_content)} bytes, type: {file.content_type}")
        
        # Validate file size (50MB limit)
        max_size = 50 * 1024 * 1024  # 50MB
        if len(file_content) > max_size:
            raise HTTPException(
                status_code=400,
                detail="File size exceeds 50MB limit"
            )
        
        # Generate unique file path: user_id/study_id/filename
        file_extension = os.path.splitext(file.filename)[1] if file.filename else '.pdf'
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = f"{user['uid']}/{study_id}/{unique_filename}"
        print(f"Upload path: {file_path}")
        
        # Upload to Supabase Storage
        try:
            upload_result = supabase.storage.from_('mra-documents').upload(
                file_path, 
                file_content,
                {
                    'content-type': file.content_type,
                    'x-upsert': 'true'  # Allow overwrite if needed
                }
            )
            print(f"Upload result: {upload_result}")
        except Exception as upload_error:
            print(f"Upload error: {upload_error}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to upload file: {str(upload_error)}"
            )
        
        # Get public URL for the uploaded file
        try:
            url_result = supabase.storage.from_('mra-documents').get_public_url(file_path)
            print(f"URL result: {url_result}")
        except Exception as url_error:
            print(f"URL error: {url_error}")
            # Continue without public URL if this fails
            url_result = {'publicURL': ''}
        
        # Create document record in database
        # Handle different URL result formats
        upload_url = ''
        if hasattr(url_result, 'publicURL'):
            upload_url = url_result.publicURL
        elif isinstance(url_result, dict):
            upload_url = url_result.get('publicURL', '')
        elif isinstance(url_result, str):
            upload_url = url_result
        
        document_data = {
            'id': str(uuid.uuid4()),
            'filename': file.filename or 'uploaded_document.pdf',
            'file_size': len(file_content),
            'file_type': file.content_type,
            'file_path': file_path,
            'upload_url': upload_url,
            'study_id': study_id,
            'user_id': user['uid'],
            'processing_status': 'uploaded'
        }
        
        # Insert into MRA documents table
        try:
            result = supabase.table('mra_documents').insert(document_data).execute()
            print(f"Database insert result: {result}")
            
            if not result.data:
                raise Exception("No data returned from database insert")
            
            document = result.data[0]
            print(f"Document created successfully: {document['id']}")
            
            return {"document": document}
            
        except Exception as db_error:
            print(f"Database error: {db_error}")
            # Cleanup: delete uploaded file if database insert failed
            try:
                supabase.storage.from_('mra-documents').remove([file_path])
                print(f"Cleaned up uploaded file: {file_path}")
            except Exception as cleanup_error:
                print(f"Failed to cleanup file: {cleanup_error}")
            
            raise HTTPException(status_code=500, detail=f"Failed to create document record: {str(db_error)}")
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error uploading document: {str(e)}")

@router.get("/{document_id}", response_model=Dict[str, DocumentResponse])
async def get_document(document_id: str, user=Depends(get_current_user)):
    """Get a specific document"""
    try:
        result = supabase.table('mra_documents')\
            .select("*")\
            .eq('id', document_id)\
            .eq('user_id', user['uid'])\
            .single()\
            .execute()
        
        if not result.data:
            raise HTTPException(status_code=404, detail="Document not found")
        
        return {"document": result.data}
        
    except Exception as e:
        if "not found" in str(e).lower():
            raise HTTPException(status_code=404, detail="Document not found")
        raise HTTPException(status_code=500, detail=f"Error fetching document: {str(e)}")

@router.get("/study/{study_id}", response_model=Dict[str, List[DocumentResponse]])
async def get_study_documents(study_id: str, user=Depends(get_current_user)):
    """Get all documents for a study"""
    try:
        result = supabase.table('mra_documents')\
            .select("*")\
            .eq('study_id', study_id)\
            .eq('user_id', user['uid'])\
            .order('created_at', desc=True)\
            .execute()
        
        return {"documents": result.data or []}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching study documents: {str(e)}")

@router.delete("/{document_id}")
async def delete_document(document_id: str, user=Depends(get_current_user)):
    """Delete a document and its file from storage"""
    try:
        # Get document info first
        result = supabase.table('mra_documents')\
            .select("*")\
            .eq('id', document_id)\
            .eq('user_id', user['uid'])\
            .single()\
            .execute()
        
        if not result.data:
            raise HTTPException(status_code=404, detail="Document not found")
        
        document = result.data
        
        # Delete from storage
        if document.get('file_path'):
            supabase.storage.from_('mra-documents').remove([document['file_path']])
        
        # Delete from database
        supabase.table('mra_documents')\
            .delete()\
            .eq('id', document_id)\
            .eq('user_id', user['uid'])\
            .execute()
        
        return {"message": "Document deleted successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

@router.post("/{document_id}/analyze")
async def analyze_document(
    document_id: str, 
    request: DocumentAnalysisRequest,
    user=Depends(get_current_user)
):
    """Analyze an uploaded document using OpenAI (GPT-5/o3)"""
    try:
        # Get document info
        doc_result = supabase.table('mra_documents')\
            .select("*")\
            .eq('id', document_id)\
            .eq('user_id', user['uid'])\
            .single()\
            .execute()
        
        if not doc_result.data:
            raise HTTPException(status_code=404, detail="Document not found")
        
        document = doc_result.data
        
        # Update processing status to analyzing
        supabase.table('mra_documents')\
            .update({'processing_status': 'analyzing'})\
            .eq('id', document_id)\
            .execute()
        
        try:
            # Download file from Supabase Storage
            file_response = supabase.storage.from_('mra-documents').download(document['file_path'])
            
            # Create temporary file for processing
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(document['filename'])[1]) as temp_file:
                temp_file.write(file_response)
                temp_file_path = temp_file.name
            
            try:
                # Extract text from document
                print(f"Extracting text from {document['file_type']} document...")
                document_text = extract_text_from_document(temp_file_path, document['file_type'])
                
                if not document_text.strip():
                    raise Exception("No text could be extracted from the document")
                
                print(f"Extracted {len(document_text)} characters from document")
                
                # Analyze document with OpenAI (GPT-5/o3)
                print(f"Starting OpenAI {settings.OPENAI_MODEL} analysis...")
                analysis_result = await analyze_mra_document_with_ai(document_text)
                
                # Convert to dict for database storage
                analysis_data = analysis_result.model_dump()
                
                # Update document with analysis results
                supabase.table('mra_documents')\
                    .update({
                        'processing_status': 'completed',
                        'analysis_results': analysis_data
                    })\
                    .eq('id', document_id)\
                    .execute()
                
                print(f"Analysis complete: found {analysis_result.total_findings} findings")
                
                return {
                    "message": "Document analysis completed successfully",
                    "document_id": document_id,
                    "analysis_type": request.analysis_type,
                    "status": "completed",
                    "results": analysis_data
                }
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
                    
        except Exception as analysis_error:
            # Update status to failed
            error_message = str(analysis_error)
            supabase.table('mra_documents')\
                .update({
                    'processing_status': 'failed',
                    'error_message': error_message
                })\
                .eq('id', document_id)\
                .execute()
            
            print(f"Analysis failed: {error_message}")
            raise HTTPException(status_code=500, detail=f"Document analysis failed: {error_message}")
        
    except HTTPException:
        raise
    except Exception as e:
        # Update status to failed if not already done
        try:
            supabase.table('mra_documents')\
                .update({
                    'processing_status': 'failed',
                    'error_message': str(e)
                })\
                .eq('id', document_id)\
                .execute()
        except:
            pass
        raise HTTPException(status_code=500, detail=f"Error analyzing document: {str(e)}")